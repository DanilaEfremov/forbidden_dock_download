import requests
from bs4 import BeautifulSoup
from docx import Document



url = input("Введите адресс вашего закрытого документа: ")

# url = 'https://docs.google.com/document/d/1thuaxyN0iSfAj1gDnHTBjZ1YzVHnQiRzsO3WBc40-XU/edit?tab=t.0#heading=h.30j0zll'
url = url[:url.find('edit')]
url = url + 'mobilebasic'

response = requests.get(url)

if response.status_code == 200:
    html_content = response.text
    with open('лекция.html', 'w') as file:
        file.write(html_content)
        print('Мы все качнулм')
else:
    print('Крииииинж')

with open('лекция.html', 'r') as file:
    html_content = file.read()


soup = BeautifulSoup(response.content, 'html.parser')

doc = Document()
doc.add_heading(soup.title.string, level=1)

for paragraph in soup.find_all('p'):
    doc.add_paragraph(paragraph.get_text())

for table in soup.find_all('table'):
    for row in table.find_all('tr'):
        cells = row.find_all(['th', 'td'])
        doc_row = doc.add_table(rows=1, cols=len(cells)).rows[0]
        for i, cell in enumerate(cells):
            doc_row.cells[i].text = cell.get_text()
    doc.add_paragraph()

doc.save('лекция.docx')
print("Готово)")